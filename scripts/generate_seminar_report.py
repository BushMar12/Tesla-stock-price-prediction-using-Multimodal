"""
Generate the Seminar 2 .docx report for 49275 Neural Networks and Fuzzy Logic.

Structure follows the marking scheme in `Seminar 2 Brief NN&FL.pdf`:
    1. Introduction (5%)
    2. Motivation, Aims and Objectives (5%)
    3. Methodology + functional block diagram (20%)
    4. Experiment and Results (15%)
    5. Analysis and Discussion (15%)
    6. Conclusion (5%)
    7. Demonstration / Contribution Statement (5%)

Run:
    python scripts/generate_seminar_report.py

Output:
    Seminar_2_Tesla_Multimodal_Report_GENERATED.docx
"""
from __future__ import annotations

import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

OUTPUT_PATH = PROJECT_ROOT / "Seminar_2_NNFL_Tesla_Multimodal_Report.docx"

PRIMARY_BLUE = RGBColor(0x1F, 0x4E, 0x9D)
ACCENT_GREY = RGBColor(0x55, 0x55, 0x55)


# --------------------------------------------------------------------------
# Style helpers
# --------------------------------------------------------------------------
def set_cell_background(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.color.rgb = PRIMARY_BLUE
        if level == 0:
            run.font.size = Pt(24)
        elif level == 1:
            run.font.size = Pt(16)
        else:
            run.font.size = Pt(13)


def add_body(doc: Document, text: str, italic: bool = False, size: int = 11) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    paragraph.paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    paragraph.paragraph_format.left_indent = Cm(0.6 + 0.6 * level)
    paragraph.paragraph_format.space_after = Pt(2)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = ACCENT_GREY
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)


def add_image(doc: Document, image_path: Path, width_inches: float = 6.0, caption: str | None = None) -> None:
    if not image_path.exists():
        add_body(doc, f"[Figure missing: {image_path.name}]", italic=True)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    if caption:
        add_caption(doc, caption)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = 1  # center
    table.style = "Light Grid Accent 1"

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr_cells[i], "1F4E9D")

    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for row in table.rows:
            for cell, width in zip(row.cells, col_widths):
                cell.width = Inches(width)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)


# --------------------------------------------------------------------------
# Data loaders
# --------------------------------------------------------------------------
def load_standalone_results() -> pd.DataFrame:
    path = PROJECT_ROOT / "models" / "model_comparison.csv"
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)


def load_multimodal_results() -> pd.DataFrame:
    path = PROJECT_ROOT / "models" / "ablation_results.csv"
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path)


def fmt_money(value: float) -> str:
    if pd.isna(value):
        return "—"
    return f"${value:.2f}"


def fmt_pct(value: float) -> str:
    if pd.isna(value):
        return "—"
    return f"{value:.2f}%"


# --------------------------------------------------------------------------
# Document sections
# --------------------------------------------------------------------------
def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Tesla Stock Price Prediction\nUsing a Multimodal Transformer Architecture")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = PRIMARY_BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Seminar 2 — Project Report")
    sub_run.font.size = Pt(14)
    sub_run.italic = True
    sub_run.font.color.rgb = ACCENT_GREY

    course = doc.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course_run = course.add_run(
        "49275 Neural Networks and Fuzzy Logic — Autumn 2026\n"
        "Faculty of Engineering and IT, University of Technology Sydney"
    )
    course_run.font.size = Pt(11)

    doc.add_paragraph()
    summary_heading = doc.add_paragraph()
    summary_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sh_run = summary_heading.add_run("Executive Summary")
    sh_run.bold = True
    sh_run.font.size = Pt(13)
    sh_run.font.color.rgb = PRIMARY_BLUE

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    summary.add_run(
        "This report compares two parallel pipelines for predicting Tesla "
        "(TSLA) next-day stock price. The first uses a custom multimodal "
        "fusion model with a Transformer-based price encoder, a temporal-"
        "CNN sentiment encoder, and explicit cross-modal attention. The "
        "second uses standalone sequence regressors (LSTM, GRU, Transformer "
        "and XGBoost) over either price-and-technical features alone or the "
        "same features concatenated with daily Alpha Vantage news sentiment. "
        "Under matched data, preprocessing and training conditions, a "
        "standalone Transformer over price-and-technical features attains "
        "the lowest test error (MAE $8.75, MAPE 2.08%), beating the best "
        "multimodal variant by approximately 9% MAE despite using "
        "substantially fewer parameters. Within the multimodal family, "
        "however, cross-attention is essential for extracting useful signal "
        "from real news sentiment. We recommend that the standalone "
        "Transformer be adopted as the production model and that the "
        "multimodal architecture be retained as a documented baseline that "
        "supports the project's negative-result contribution."
    )
    for run in summary.runs:
        run.font.size = Pt(11)
    doc.add_page_break()


def add_introduction(doc: Document) -> None:
    add_heading(doc, "1. Introduction", level=1)
    add_body(
        doc,
        "Predicting short-term equity prices is a long-standing problem in "
        "computational finance with two well-known difficulties. First, "
        "daily price changes are dominated by noise — the signal-to-noise "
        "ratio of any single feature is low. Second, prices respond not "
        "only to historical price movements but also to qualitative "
        "information that is hard to encode numerically, such as news, "
        "macroeconomic context and investor sentiment.",
    )
    add_body(
        doc,
        "This project targets the next-day closing price of Tesla, Inc. "
        "(TSLA) — a stock that is widely held, heavily traded, and known "
        "for sharp idiosyncratic moves driven by news. We treat the daily "
        "prediction problem as a sequence regression task: given a 60-day "
        "window of price-derived and sentiment-derived features, predict "
        "the next-day return, then reconstruct the next-day closing price.",
    )
    add_body(
        doc,
        "Two architectural families are compared. The first is a "
        "purpose-built multimodal model that keeps the price stream and the "
        "sentiment stream in separate encoders and fuses them through a "
        "cross-modal attention block. The second is a family of standalone "
        "sequence regressors (LSTM, GRU, Transformer, XGBoost) that operate "
        "on a single tensor formed by concatenating the two streams along "
        "the feature axis. Both pipelines share the same preprocessing, "
        "splits, scalers and evaluation metrics, so any difference in test "
        "error can be attributed to the architecture rather than to data "
        "leakage or evaluation drift.",
    )


def add_motivation(doc: Document) -> None:
    add_heading(doc, "2. Motivation, Aims and Objectives", level=1)

    add_heading(doc, "2.1 Recap of Seminar 1", level=2)
    add_body(
        doc,
        "In Seminar 1 we proposed a multimodal Transformer that would jointly "
        "model price/technical signals and news-derived sentiment for TSLA. "
        "The hypothesis was that explicit cross-modal fusion — keeping each "
        "modality in its own encoder and connecting them through attention — "
        "would outperform single-stream models that flatten the modalities "
        "into a single tensor. The expected outcome was that adding sentiment, "
        "and adding cross-attention on top of sentiment, would each "
        "monotonically reduce test MAE on next-day TSLA price.",
    )

    add_heading(doc, "2.2 Aims", level=2)
    add_bullet(doc, "Build an end-to-end pipeline that ingests TSLA price and Alpha Vantage news data, computes technical indicators and daily sentiment features, and trains a multimodal sequence model end-to-end.")
    add_bullet(doc, "Quantify the contribution of (a) sentiment and (b) cross-modal attention through a controlled ablation study.")
    add_bullet(doc, "Compare the multimodal model against canonical sequence-regression baselines (LSTM, GRU, Transformer, XGBoost) under matched conditions.")
    add_bullet(doc, "Deliver an inference-ready system (Streamlit dashboard with SHAP-based explainability) that exposes predictions and feature attributions to the end user.")

    add_heading(doc, "2.3 Objectives", level=2)
    add_bullet(doc, "Implement DataPreprocessor to merge price, technical and sentiment streams, fit scalers on the training portion only, and emit chronological 60-day sliding windows.")
    add_bullet(doc, "Implement MultimodalFusionModel with a Transformer price encoder, a temporal-CNN sentiment encoder and a price-to-sentiment cross-attention block.")
    add_bullet(doc, "Implement four baseline architectures (LSTM, GRU, Transformer, XGBoost) sharing the same preprocessor.")
    add_bullet(doc, "Train all models with the same seed, optimizer, schedule and metric, and evaluate on the same chronological test split.")

    add_heading(doc, "2.4 What changed since Seminar 1", level=2)
    add_body(
        doc,
        "Two findings during implementation forced a substantial revision "
        "of the project's expected outcome:",
    )
    add_bullet(doc, "A bug in the standalone pipeline (sentiment loaded from a generic CSV path rather than routed through fetch_sentiment_data) meant that the original benchmark was using synthetic sentiment instead of real Alpha Vantage news. The fix made the standalone pipeline a fair real-sentiment competitor for the first time.")
    add_bullet(doc, "Once the bug was fixed, the standalone Transformer outperformed the multimodal model. The project's contribution therefore shifted from \"propose a multimodal architecture\" to \"empirically test whether cross-modal fusion outperforms a single-stream Transformer on this problem,\" with the answer being \"no, on this dataset.\"")


def add_methodology(doc: Document) -> None:
    add_heading(doc, "3. Methodology", level=1)

    add_heading(doc, "3.1 Functional block diagram", level=2)
    block_diagram = PROJECT_ROOT / "pipeline_architecture.png"
    add_image(
        doc,
        block_diagram,
        width_inches=6.3,
        caption="Figure 1. End-to-end pipeline. Two aligned daily streams (price/technical and sentiment) are date-merged, scaled, windowed and fed to either a multimodal fusion model or a standalone sequence regressor. Both pipelines share preprocessing, splits and evaluation.",
    )

    add_heading(doc, "3.2 Data sources and dataset structure", level=2)
    add_bullet(doc, "Stock data — Yahoo Finance OHLCV for TSLA, 2021-01-01 to 2026-04-21 (≈1,300 trading days). Cached to data/raw/TSLA_historical.csv.")
    add_bullet(doc, "Sentiment data — Alpha Vantage NEWS_SENTIMENT API. Per-article TSLA-specific sentiment scores aggregated to daily means; cached to data/raw/alpha_vantage_tsla_sentiment_data.csv (1,381 daily rows).")
    add_bullet(doc, "Market context — SPY and VIX OHLCV from Yahoo Finance, currently disabled in config but cached for future use.")
    add_body(
        doc,
        "The training tensor for each window contains 10 price/technical features "
        "(open, high, low, close, volume, SMA-20, SMA-50, RSI-14, Bollinger Upper, "
        "Middle, Lower) and 9 sentiment features (compound, positive, negative, "
        "neutral, news_count, plus 1/2/3-day lagged compound). Each sample is a "
        "(60, 10) or (60, 19) tensor depending on whether sentiment is included.",
    )

    add_heading(doc, "3.3 Preprocessing", level=2)
    add_bullet(doc, "Date-merge stock, technical and sentiment tables on trading dates. Missing-news days fill with neutral sentiment (compound=0, neutral=1).")
    add_bullet(doc, "Drop infinities and forward-fill the small number of missing technical values produced by rolling-window indicators.")
    add_bullet(doc, "Fit MinMax / Standard scalers on the training portion only to prevent test-set leakage.")
    add_bullet(doc, "Construct overlapping 60-day windows. Target for each window is the next-day return, scaled by the same return scaler used at inference.")
    add_bullet(doc, "Chronological 80/10/10 train/validation/test split (no shuffling). Test set spans the most recent ~128 trading days.")

    add_heading(doc, "3.4 Models compared", level=2)
    add_table(
        doc,
        ["Pipeline", "Model / Mode", "Sentiment", "Cross-attention", "Params"],
        [
            ["Standalone", "LSTM", "off / concat", "—", "small"],
            ["Standalone", "GRU", "off / concat", "—", "small"],
            ["Standalone", "Transformer", "off / concat", "self-attn over concat", "small"],
            ["Standalone", "XGBoost", "off / concat", "—", "n/a"],
            ["Multimodal", "no-sentiment", "off", "off", "1.20M"],
            ["Multimodal", "sentiment, no cross-attn", "on, MLP fusion", "off", "1.34M"],
            ["Multimodal", "current (sentiment + cross-attn)", "on, CNN encoder", "on", "1.67M"],
        ],
        col_widths=[1.0, 2.4, 1.5, 1.6, 0.8],
    )

    add_heading(doc, "3.5 Training configuration", level=2)
    add_table(
        doc,
        ["Setting", "Value"],
        [
            ["Sequence length", "60 days"],
            ["Batch size", "32"],
            ["Optimiser", "AdamW, learning rate 1e-4"],
            ["LR schedule", "ReduceLROnPlateau, factor 0.5, patience 5 (val MAE in dollar space)"],
            ["Loss", "Smooth L1 on scaled return"],
            ["Maximum epochs", "50"],
            ["Random seed", "42"],
            ["Determinism", "torch.use_deterministic_algorithms(True), cuDNN deterministic"],
            ["Sentiment source", "alpha_vantage (cached daily aggregation)"],
            ["Train / val / test split", "80% / 10% / 10% chronological"],
        ],
        col_widths=[2.0, 4.5],
    )
    add_body(
        doc,
        "All metrics are reported in dollar price space. Predicted scaled "
        "returns are inverse-transformed with the fitted return scaler and "
        "reconstructed into next-day prices via close_t × (1 + return_pred) "
        "before any error metric is computed.",
    )


def add_results(doc: Document, standalone: pd.DataFrame, multimodal: pd.DataFrame) -> None:
    add_heading(doc, "4. Experiment and Results", level=1)

    add_heading(doc, "4.1 Standalone pipeline (single seed)", level=2)
    if not standalone.empty:
        rows = []
        for _, r in standalone.sort_values("MAE").iterrows():
            rows.append([
                r["Configuration"],
                r["Model"],
                fmt_money(r["RMSE"]),
                fmt_money(r["MAE"]),
                fmt_pct(r["MAPE"]),
            ])
        add_table(
            doc,
            ["Configuration", "Model", "RMSE", "MAE", "MAPE"],
            rows,
            col_widths=[1.7, 1.2, 1.0, 1.0, 1.0],
        )
        add_caption(doc, "Table 1. Standalone-pipeline test-set metrics (seed=42, 50 epochs, real Alpha Vantage sentiment), sorted by MAE.")

    add_heading(doc, "4.2 Multimodal ablation (single seed)", level=2)
    if not multimodal.empty:
        rows = []
        for _, r in multimodal.sort_values("mae").iterrows():
            rows.append([
                r["mode"],
                r["sentiment_source"],
                f"{int(r['params']):,}",
                fmt_money(r["rmse"]),
                fmt_money(r["mae"]),
                fmt_pct(r["mape"]),
                str(int(r["best_epoch"])),
            ])
        add_table(
            doc,
            ["Mode", "Sentiment", "Params", "RMSE", "MAE", "MAPE", "Best epoch"],
            rows,
            col_widths=[1.6, 1.0, 0.9, 0.8, 0.8, 0.8, 0.7],
        )
        add_caption(doc, "Table 2. Multimodal-ablation test-set metrics (real Alpha Vantage sentiment, seed=42, 50 epochs), sorted by MAE.")

    ablation_png = PROJECT_ROOT / "models" / "ablation_comparison.png"
    if ablation_png.exists():
        add_image(
            doc,
            ablation_png,
            width_inches=6.3,
            caption="Figure 2. Multimodal-ablation test metrics (RMSE / MAE / MAPE) by mode, with standalone Transformer and XGBoost baselines overlaid as dashed lines.",
        )

    curves_png = PROJECT_ROOT / "models" / "ablation_train_val_curves.png"
    if curves_png.exists():
        add_image(
            doc,
            curves_png,
            width_inches=6.3,
            caption="Figure 3. Train / validation loss curves for the three multimodal modes. The 'sentiment, no cross-attention' mode reaches its best validation loss at epoch 1 and degrades thereafter — evidence that real sentiment cannot be productively fused without attention.",
        )

    add_heading(doc, "4.3 Combined ranking", level=2)
    add_body(
        doc,
        "Ranking all variants from both pipelines under matched seed, "
        "epochs and sentiment-source conditions:",
    )
    combined_rows = []
    if not standalone.empty:
        for _, r in standalone.iterrows():
            combined_rows.append({
                "pipeline": "Standalone",
                "variant": f"{r['Model']} ({r['Configuration']})",
                "MAE": float(r["MAE"]),
                "RMSE": float(r["RMSE"]),
                "MAPE": float(r["MAPE"]),
            })
    if not multimodal.empty:
        for _, r in multimodal.iterrows():
            combined_rows.append({
                "pipeline": "Multimodal",
                "variant": f"fusion ({r['mode']})",
                "MAE": float(r["mae"]),
                "RMSE": float(r["rmse"]),
                "MAPE": float(r["mape"]),
            })
    if combined_rows:
        combined_df = pd.DataFrame(combined_rows).sort_values("MAE")
        rows = []
        for rank, (_, r) in enumerate(combined_df.iterrows(), start=1):
            rows.append([str(rank), r["pipeline"], r["variant"], fmt_money(r["MAE"]), fmt_money(r["RMSE"]), fmt_pct(r["MAPE"])])
        add_table(
            doc,
            ["Rank", "Pipeline", "Variant", "MAE", "RMSE", "MAPE"],
            rows,
            col_widths=[0.55, 1.0, 2.5, 0.8, 0.8, 0.8],
        )
        add_caption(doc, "Table 3. Combined ranking of all 11 model variants (single-seed, real Alpha Vantage sentiment).")


def add_analysis(doc: Document) -> None:
    add_heading(doc, "5. Analysis and Discussion", level=1)

    add_heading(doc, "5.1 Why does the standalone Transformer beat the multimodal model?", level=2)
    add_body(
        doc,
        "The multimodal model carries 1.67M parameters yet generalises worse "
        "than a much smaller standalone Transformer over concatenated features. "
        "Three factors plausibly explain this gap:",
    )
    add_bullet(doc, "Capacity-to-data ratio. With ~1,013 training windows, a 1.67M-parameter network has high capacity per training sample. The mild train/val gap (-0.16%) suggests the multimodal model is not catastrophically overfitting, but the extra parameters do not translate into lower test loss.")
    add_bullet(doc, "Forced separation imposes a weak inductive bias. The multimodal architecture commits to \"price and sentiment are different modalities, encode them separately and fuse once.\" A single Transformer over concat(price, sentiment) makes no such commitment — its self-attention can route across price and sentiment dimensions at every layer. When cross-modal interactions are simple, separation costs more than it gains.")
    add_bullet(doc, "Sentiment encoder bottleneck. The temporal-CNN sentiment encoder compresses the 9-dimensional sentiment sequence before fusion. If the price branch could have used the lag-1 sentiment compound score directly, the encoder may have already discarded that information. A single Transformer over the concatenated tensor sidesteps the bottleneck — every dimension is available to every attention layer.")
    add_body(
        doc,
        "This is consistent with broader findings in deep-learning architecture "
        "research: explicit modality separation is most useful when the "
        "modalities have very different statistical structure or sampling rates "
        "(e.g., text + image, asynchronous video + audio). For two daily-aligned "
        "numerical streams the inductive bias is too weak to recoup the parameter cost.",
    )

    add_heading(doc, "5.2 Sentiment is useful, but only with attention", level=2)
    add_body(
        doc,
        "A second consistent thread runs across both pipelines: real news "
        "sentiment improves test error only when paired with an attention "
        "mechanism that can filter it.",
    )
    add_bullet(doc, "Recurrent baselines (LSTM, GRU) without attention degrade when sentiment is concatenated to the input — every byte of sentiment dilutes the price signal in the recurrent state.")
    add_bullet(doc, "The multimodal \"sentiment, no cross-attention\" mode degrades worst of all (best epoch 1, then collapses). Its MLP fusion has even less filtering capacity than a recurrent state.")
    add_bullet(doc, "The Transformer (in either pipeline) and the multimodal cross-attention mode benefit from real sentiment, or at least do not degrade.")
    add_body(
        doc,
        "This matters because earlier ablation results, run on synthetic "
        "sentiment, had picked the simpler MLP-fusion variant as the multimodal "
        "winner. Synthetic sentiment was a smoothed function of lagged returns — "
        "easy to fuse linearly. Real news sentiment is genuinely noisy and "
        "non-linearly related to next-day return, which only attention can handle. "
        "For noisy auxiliary modalities, attention is not optional; it appears necessary.",
    )

    add_heading(doc, "5.3 The XGBoost regression", level=2)
    add_body(
        doc,
        "XGBoost is consistently the worst architecture in both configurations. "
        "Its native early-stopping criterion picks a very early best round (≤8 "
        "boosting rounds for technical, 1 for technical+sentiment); validation "
        "MAE rises monotonically thereafter. The mechanical cause is clear — "
        "XGBoost flattens the 60-day window into a fixed-length feature vector, "
        "discarding the temporal ordering that the sequence models exploit. We "
        "retain XGBoost in the report as a useful \"non-temporal baseline\" "
        "reference point but not as a competitive option.",
    )

    add_heading(doc, "5.4 SHAP-based feature attribution (sanity check)", level=2)
    shap_png = PROJECT_ROOT / "models" / "shap_feature_importance.png"
    if shap_png.exists():
        add_image(
            doc,
            shap_png,
            width_inches=5.5,
            caption="Figure 4. SHAP global feature importance from an XGBoost surrogate trained on the same windows, used as a sanity check on what the standalone pipeline learns.",
        )
    add_body(
        doc,
        "The SHAP attribution highlights short-window price-derived features "
        "(close, recent moving averages, RSI) as the dominant drivers, with "
        "lagged sentiment compound scores contributing a smaller but non-zero "
        "share of attribution. This is consistent with the head-to-head result: "
        "sentiment carries some signal but its marginal contribution is small "
        "and depends on architecture.",
    )

    add_heading(doc, "5.5 Optimisation strategies considered", level=2)
    add_bullet(doc, "Reduced multimodal capacity (fewer transformer layers, narrower hidden states) — would test whether the multimodal gap is purely a capacity-vs-data issue.")
    add_bullet(doc, "Replacing the temporal-CNN sentiment encoder with a Transformer or a lightweight identity-with-positional-encoding pass-through.")
    add_bullet(doc, "Walk-forward expanding-window evaluation rather than a single chronological 80/10/10 split, to test ranking stability across market regimes.")
    add_bullet(doc, "Direction-loss training (binary or multiclass) rather than smooth L1 on return — already partly explored in earlier project iterations and noted in commit f9c6ee5.")
    add_bullet(doc, "Higher-frequency or relevance-weighted sentiment aggregation (weight by Alpha Vantage relevance score rather than uniform mean).")


def add_conclusion(doc: Document) -> None:
    add_heading(doc, "6. Conclusion", level=1)
    add_body(
        doc,
        "We set out to test whether explicit cross-modal fusion via a "
        "purpose-built multimodal architecture would outperform single-stream "
        "sequence models on next-day TSLA price prediction. Under matched "
        "data, preprocessing and training conditions, the answer is no: a "
        "standalone Transformer over concatenated price-and-technical features "
        "achieves the lowest test MAE ($8.75; MAPE 2.08%), beating the best "
        "multimodal variant (cross-attention enabled, MAE $9.53) by "
        "approximately 9% while using substantially fewer parameters. Within "
        "the multimodal family, however, cross-attention is essential — the "
        "simpler MLP-fusion baseline collapses on real news sentiment.",
    )
    add_body(
        doc,
        "These findings reframe the project's contribution. Rather than "
        "presenting a multimodal architecture as the recommended solution, the "
        "report contributes a controlled comparison: for two daily-aligned "
        "numerical streams (price and sentiment) on a single equity, a single "
        "self-attention model over concatenated features is sufficient and "
        "preferable to forced modality separation. The multimodal codebase is "
        "retained as a documented baseline so that the negative result remains "
        "reproducible from the same repository.",
    )

def add_demo_and_contribution(doc: Document) -> None:
    add_heading(doc, "7. Demonstration and Contribution Statement", level=1)

    add_heading(doc, "7.1 Demonstration", level=2)
    add_body(
        doc,
        "The accompanying one-minute demonstration video shows: (i) launching the "
        "Streamlit dashboard at app/streamlit_app.py, (ii) viewing the next-day "
        "price prediction card and price-overlay chart, (iii) the model-comparison "
        "tab listing the four standalone baselines and their test metrics, and "
        "(iv) the SHAP explainability tab showing the global feature-importance "
        "bar chart and the temporal SHAP heatmap. The dashboard runs the standalone "
        "Transformer (technical+sentiment) checkpoint by default; the multimodal "
        "checkpoint is loadable as an ablation comparison.",
    )
    add_body(
        doc,
        "Reproduction commands for the headline results:",
    )
    add_bullet(doc, "python model_comparison.py --seed 42 --epochs 50  # standalone, single seed")
    add_bullet(doc, "python scripts/run_ablation_validation.py --sentiment-source alpha_vantage --seed 42 --epochs 50  # multimodal ablation")
    add_bullet(doc, "streamlit run app/streamlit_app.py  # dashboard")

    add_heading(doc, "7.2 Contribution Statement", level=2)
    add_body(
        doc,
        "All authors contributed jointly to the project design, the "
        "interpretation of results and the writing of this report. Specific "
        "responsibilities are summarised below; please replace the placeholders "
        "with member names before submission.",
    )
    add_table(
        doc,
        ["Member", "Primary contributions"],
        [
            ["[Member 1]", "Data ingestion (stock_data.py, sentiment_data.py — Yahoo Finance + Alpha Vantage), preprocessing pipeline (preprocessing.py), DataPreprocessor sliding-window logic, sentiment-loading bug fix."],
            ["[Member 2]", "Multimodal model (fusion.py — Transformer price encoder, temporal-CNN sentiment encoder, cross-modal attention block), trainer (trainer.py), ablation runner (run_ablation_validation.py)."],
            ["[Member 3]", "Standalone baselines (regression_models.py — LSTM, GRU, Transformer, XGBoost), model_comparison.py training driver, evaluation utilities, comparison plots."],
            ["[Member 4]", "Streamlit dashboard (app/streamlit_app.py), SHAP explainability (shap_feature_importance.py), inference script (predict.py), figure preparation and report drafting."],
        ],
        col_widths=[1.5, 5.0],
    )
    add_body(
        doc,
        "The automated docx report generator (scripts/generate_seminar_report.py) "
        "was jointly produced and is available in the repository for reproducibility.",
        italic=True,
    )


def main() -> None:
    standalone = load_standalone_results()
    multimodal = load_multimodal_results()

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    add_cover(doc)
    add_introduction(doc)
    add_motivation(doc)
    add_methodology(doc)
    add_results(doc, standalone, multimodal)
    add_analysis(doc)
    add_conclusion(doc)
    add_demo_and_contribution(doc)

    doc.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")
    print(f"  Standalone rows: {len(standalone)}")
    print(f"  Multimodal rows: {len(multimodal)}")


if __name__ == "__main__":
    main()
